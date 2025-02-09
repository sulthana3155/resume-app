from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
import os

app = Flask(__name__)

# Folder for storing generated resumes and uploaded files
RESUME_FOLDER = 'generated_resumes'
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(RESUME_FOLDER):
    os.makedirs(RESUME_FOLDER)
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Job suggestion API based on skills
def get_job_suggestions(skills):
    # Suggest jobs based on skills
    job_suggestions = []
    if "Python" in skills:
        job_suggestions.extend(["Google", "Facebook", "Microsoft", "Netflix"])
    if "Java" in skills:
        job_suggestions.extend(["Oracle", "IBM", "Amazon", "Spotify"])
    if "Machine Learning" in skills:
        job_suggestions.extend(["Tesla", "NVIDIA", "DeepMind", "OpenAI"])
    if "Web Development" in skills:
        job_suggestions.extend(["Trello", "GitHub", "Atlassian", "Spotify"])
    
    if not job_suggestions:
        job_suggestions.append("Startup Inc.")  # Default job suggestion for unknown skills

    return job_suggestions

# Function to extract text from uploaded DOCX file
def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        return str(e)

@app.route('/', methods=['GET', 'POST'])
def home():
    job_suggestions = []

    if request.method == 'POST':
        # Get user inputs from the form
        name = request.form['name']
        email = request.form['email']
        experience = request.form['experience']
        education = request.form['education']
        skills = request.form['skills'].split(',')  # Assuming skills are comma-separated

        # Get the uploaded resume file (optional)
        uploaded_file = request.files.get('upload_resume')
        uploaded_filename = None
        extracted_text = ""

        if uploaded_file:
            # Save the uploaded file to the server
            uploaded_filename = os.path.join(UPLOAD_FOLDER, "resume_template.docx")
            uploaded_file.save(uploaded_filename)
            
            # Extract text from the uploaded DOCX file
            extracted_text = extract_text_from_docx(uploaded_filename)

            # Optionally, extract skills and other info from the document text (if needed)
            # Here, we're just adding the extracted skills from the document
            extracted_skills = ["Python" if "Python" in extracted_text else None, 
                                "Java" if "Java" in extracted_text else None]
            skills.extend([skill for skill in extracted_skills if skill])

        # Ensure all required fields are provided
        if not all([name, email, experience, education, skills]):
            return jsonify({"error": "All fields are required!"}), 400

        # Job suggestions based on skills
        job_suggestions = get_job_suggestions(skills)

        # Generate the resume based on a template
        template_path = './uploads/resume_template.docx'  # Path to the Word template
        try:
            doc = Document(template_path)
        except Exception as e:
            return jsonify({"error": f"Error opening template: {str(e)}"}), 500

        # Replace placeholders with user input in the document
        for para in doc.paragraphs:
            if '<<NAME>>' in para.text:
                para.text = para.text.replace('<<NAME>>', name)
            if '<<EMAIL>>' in para.text:
                para.text = para.text.replace('<<EMAIL>>', email)
            if '<<EXPERIENCE>>' in para.text:
                para.text = para.text.replace('<<EXPERIENCE>>', experience)
            if '<<EDUCATION>>' in para.text:
                para.text = para.text.replace('<<EDUCATION>>', education)
            if '<<SKILLS>>' in para.text:
                para.text = para.text.replace('<<SKILLS>>', ', '.join(skills))

        # Save the generated resume
        resume_path = os.path.join(RESUME_FOLDER, f"{name.replace(' ', '_')}_resume.docx")
        doc.save(resume_path)

        # Return the generated resume for download
        return render_template('index.html', job_suggestions=job_suggestions)

    return render_template('index.html', job_suggestions=job_suggestions)

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5000)
